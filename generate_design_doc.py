"""DESIGN.md の内容を Word (.docx) ファイルとして出力するスクリプト"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


def set_cell_shading(cell, color_hex):
    """セルの背景色を設定"""
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color_hex,
        qn('w:val'): 'clear',
    })
    shading.append(shading_elem)


def add_table(doc, headers, rows):
    """テーブルを追加"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # ヘッダー行
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
        set_cell_shading(cell, '2F5496')
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)

    # データ行
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
            if r_idx % 2 == 1:
                set_cell_shading(cell, 'D6E4F0')

    doc.add_paragraph()


def add_code_block(doc, text):
    """コードブロックを追加"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(30, 30, 30)


def main():
    doc = Document()

    # デフォルトフォント設定
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Yu Gothic'
    font.size = Pt(10.5)

    # ===== 表紙 =====
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('DQ1-Sample 設計書')
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)

    doc.add_paragraph()

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('ファミコン版ドラゴンクエスト1 再現プロジェクト')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(100, 100, 100)

    for _ in range(4):
        doc.add_paragraph()

    info_items = [
        ('プロジェクト名', 'DQ1-Sample'),
        ('言語 / フレームワーク', 'Python 3.11+ / Pygame 2.5+'),
        ('作成日', '2026-02-17'),
    ]
    for label, value in info_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'{label}: ')
        run.bold = True
        run.font.size = Pt(11)
        run = p.add_run(value)
        run.font.size = Pt(11)

    doc.add_page_break()

    # ===== 目次 =====
    doc.add_heading('目次', level=1)
    toc_items = [
        '1. システム概要',
        '2. アーキテクチャ設計',
        '3. モジュール構成',
        '4. 画面・定数設計',
        '5. ステート管理',
        '6. マップシステム',
        '7. プレイヤーシステム',
        '8. 戦闘システム',
        '9. NPC・ショップ・宿屋システム',
        '10. UIシステム',
        '11. 描画システム',
        '12. データ定義',
        '13. 既知の制限・今後の課題',
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)

    doc.add_page_break()

    # ===== 1. システム概要 =====
    doc.add_heading('1. システム概要', level=1)

    doc.add_heading('1.1 プロジェクト概要', level=2)
    doc.add_paragraph(
        '本プロジェクトはファミコン版「ドラゴンクエスト1」を Python + Pygame で再現する2D RPGです。'
        'タイルベースのマップ移動、ターン制1対1戦闘、NPC会話、装備・アイテム・呪文システムを備えています。'
    )

    doc.add_heading('1.2 技術スタック', level=2)
    add_table(doc,
        ['項目', '内容'],
        [
            ['言語', 'Python 3.11+'],
            ['フレームワーク', 'Pygame 2.5+'],
            ['描画方式', 'プログラム描画（スプライトアセット不使用）'],
            ['フォント', '日本語対応システムフォント自動検出'],
            ['マップ形式', '文字列配列ベース'],
        ]
    )

    doc.add_heading('1.3 操作方法', level=2)
    add_table(doc,
        ['キー', '操作'],
        [
            ['方向キー', '移動'],
            ['Z', '決定 / はなす'],
            ['X', 'キャンセル'],
            ['Space', 'メニュー'],
            ['Escape', '終了'],
        ]
    )

    # ===== 2. アーキテクチャ設計 =====
    doc.add_heading('2. アーキテクチャ設計', level=1)

    doc.add_heading('2.1 全体構成図', level=2)
    add_code_block(doc,
        'main.py（エントリポイント）\n'
        '  └── Game（メインループ管理）\n'
        '        ├── state_stack[]  ← ステートスタックパターン\n'
        '        │   ├── TitleState      （タイトル画面）\n'
        '        │   ├── FieldState      （フィールド探索）\n'
        '        │   └── BattleState     （戦闘画面）\n'
        '        ├── Player              （プレイヤーデータ）\n'
        '        └── pygame.display      （描画サーフェス）'
    )

    doc.add_heading('2.2 設計パターン', level=2)
    add_table(doc,
        ['パターン', '適用箇所', '説明'],
        [
            ['State Stack', 'Game.state_stack', 'ステートを積み上げ/取り除きで画面遷移を管理'],
            ['Message Queue', 'MessageWindow', 'メッセージを順次表示するキュー管理'],
            ['Weighted Random', 'エンカウントテーブル', '敵の出現率を重み付きで制御'],
            ['Font Cache', 'Font クラス', 'フォントをシングルトンでキャッシュ'],
            ['Tile Render Cache', 'TileMap._render_tile', 'タイル画像を初回描画時にキャッシュ'],
        ]
    )

    doc.add_heading('2.3 ステート遷移図', level=2)
    add_code_block(doc,
        '┌────────────┐    Zキー     ┌────────────┐   エンカウント   ┌────────────┐\n'
        '│ TitleState │ ──────────→ │ FieldState │ ─────────────→ │BattleState │\n'
        '│ (タイトル)  │  change      │ (フィールド) │    push         │  (戦闘)    │\n'
        '└────────────┘             └────────────┘ ←───────────── └────────────┘\n'
        '                                │    ↑         pop\n'
        '                                │    │\n'
        '                           load_map (ワープ)\n'
        '                           マップ切替は同一ステート内で処理'
    )

    doc.add_page_break()

    # ===== 3. モジュール構成 =====
    doc.add_heading('3. モジュール構成', level=1)

    doc.add_heading('3.1 ファイル一覧', level=2)
    add_code_block(doc,
        'DQ1-Sample/\n'
        '├── main.py              # エントリポイント\n'
        '├── requirements.txt     # 依存パッケージ (pygame)\n'
        '└── src/\n'
        '    ├── __init__.py      # パッケージ初期化（空）\n'
        '    ├── config.py        # 定数・設定\n'
        '    ├── data.py          # ゲームデータ定義\n'
        '    ├── player.py        # プレイヤークラス\n'
        '    ├── ui.py            # UI コンポーネント群\n'
        '    ├── tilemap.py       # タイルマップ描画・当たり判定\n'
        '    ├── game.py          # メインゲームループ・ステート管理\n'
        '    ├── title.py         # タイトル画面ステート\n'
        '    ├── field.py         # フィールド探索ステート\n'
        '    └── battle.py        # 戦闘ステート'
    )

    doc.add_heading('3.2 モジュール依存関係', level=2)
    add_code_block(doc,
        'main.py\n'
        '  └── game.py\n'
        '        ├── config.py\n'
        '        ├── player.py ← data.py (LEVEL_TABLE, SPELLS, ITEMS, WEAPONS, ARMORS, SHIELDS)\n'
        '        ├── title.py ← config.py, ui.py\n'
        '        ├── field.py ← config.py, data.py, player.py, ui.py, tilemap.py, battle.py\n'
        '        └── battle.py ← config.py, data.py, player.py, ui.py'
    )

    doc.add_heading('3.3 各モジュールの責務', level=2)
    add_table(doc,
        ['モジュール', '責務'],
        [
            ['main.py', 'アプリケーション起動。Gameインスタンスを生成し実行'],
            ['config.py', 'タイルサイズ、画面サイズ、色、キー設定、タイル定義等の定数管理'],
            ['data.py', '敵、アイテム、呪文、レベルテーブル、マップ、NPC等の全ゲームデータ'],
            ['player.py', 'プレイヤーのステータス、装備、インベントリ、戦闘計算式'],
            ['ui.py', 'DQスタイルのウィンドウ、テキスト表示、メニュー、ステータス表示'],
            ['tilemap.py', 'マップのタイル描画、カメラ処理、当たり判定、キャラクター描画'],
            ['game.py', 'メインループ、イベント処理、ステートスタック管理'],
            ['title.py', 'タイトル画面の表示と入力処理'],
            ['field.py', 'マップ移動、NPC会話、ショップ、エンカウント、ワープ処理'],
            ['battle.py', 'ターン制戦闘の全処理（コマンド選択〜勝敗判定）'],
        ]
    )

    doc.add_page_break()

    # ===== 4. 画面・定数設計 =====
    doc.add_heading('4. 画面・定数設計', level=1)

    doc.add_heading('4.1 画面設定', level=2)
    add_table(doc,
        ['定数', '値', '説明'],
        [
            ['TILE_SIZE', '48px', '1タイルのピクセルサイズ'],
            ['SCREEN_TILES_X', '16', '横方向の表示タイル数'],
            ['SCREEN_TILES_Y', '15', '縦方向の表示タイル数'],
            ['SCREEN_WIDTH', '768px', '画面幅 (16 × 48)'],
            ['SCREEN_HEIGHT', '720px', '画面高さ (15 × 48)'],
            ['FPS', '60', 'フレームレート'],
            ['MOVE_FRAMES', '8', '1タイル移動のフレーム数'],
        ]
    )

    doc.add_heading('4.2 カラーパレット（NESスタイル）', level=2)
    add_table(doc,
        ['定数', 'RGB値', '用途'],
        [
            ['BLACK', '(0, 0, 0)', '背景色'],
            ['WHITE', '(252, 252, 252)', '文字色・枠線'],
            ['WINDOW_BG', '(0, 0, 0)', 'ウィンドウ背景'],
            ['WINDOW_BORDER', '(252, 252, 252)', 'ウィンドウ枠'],
        ]
    )

    doc.add_heading('4.3 タイル種別一覧', level=2)
    add_table(doc,
        ['定数', 'ID', '文字', '色 (RGB)', '通行', 'エンカウント'],
        [
            ['TILE_GRASS', '0', 'G', '(34,139,34)', '可', 'あり'],
            ['TILE_WATER', '1', 'W', '(65,105,225)', '不可', '-'],
            ['TILE_MOUNTAIN', '2', 'M', '(139,90,43)', '不可', '-'],
            ['TILE_FOREST', '3', 'F', '(0,100,0)', '可', 'あり'],
            ['TILE_DESERT', '4', 'D', '(210,180,140)', '可', 'あり'],
            ['TILE_BRIDGE', '5', 'B', '(139,90,43)', '可', 'あり'],
            ['TILE_FLOOR', '6', '.', '(139,119,101)', '可', '-'],
            ['TILE_WALL', '7', '#', '(105,105,105)', '不可', '-'],
            ['TILE_DOOR', '8', '+', '(139,90,43)', '可', '-'],
            ['TILE_STAIRS_DOWN', '9', '>', '(80,80,80)', '可', '-'],
            ['TILE_STAIRS_UP', '10', '<', '(80,80,80)', '可', '-'],
            ['TILE_CASTLE', '11', 'C', '(192,192,192)', '可', '-'],
            ['TILE_TOWN', '12', 'T', '(160,82,45)', '可', '-'],
            ['TILE_CAVE', '13', 'V', '(64,64,64)', '可', '-'],
            ['TILE_SWAMP', '14', 'S', '(75,0,130)', '可', 'あり'],
            ['TILE_ROOF', '15', 'R', '(178,34,34)', '不可', '-'],
            ['TILE_COUNTER', '16', 'K', '(139,90,43)', '不可', '-'],
            ['TILE_CHEST', '17', 'X', '(218,165,32)', '可', '-'],
            ['TILE_THRONE', '18', 'O', '(218,165,32)', '可', '-'],
        ]
    )

    doc.add_heading('4.4 キー設定', level=2)
    add_table(doc,
        ['定数', 'Pygameキー', '操作'],
        [
            ['KEY_UP', 'K_UP', '上移動'],
            ['KEY_DOWN', 'K_DOWN', '下移動'],
            ['KEY_LEFT', 'K_LEFT', '左移動'],
            ['KEY_RIGHT', 'K_RIGHT', '右移動'],
            ['KEY_CONFIRM', 'K_z', '決定 / はなす'],
            ['KEY_CANCEL', 'K_x', 'キャンセル'],
            ['KEY_MENU', 'K_SPACE', 'メニュー開閉'],
        ]
    )

    doc.add_page_break()

    # ===== 5. ステート管理 =====
    doc.add_heading('5. ステート管理', level=1)

    doc.add_heading('5.1 Game クラス', level=2)
    add_code_block(doc,
        'class Game:\n'
        '    screen: pygame.Surface      # 描画サーフェス\n'
        '    clock: pygame.time.Clock    # FPS制御\n'
        '    running: bool               # ゲームループフラグ\n'
        '    player: Player              # プレイヤーインスタンス\n'
        '    state_stack: list           # ステートスタック'
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('ステート操作メソッド:')
    run.bold = True

    add_table(doc,
        ['メソッド', '動作', '使用例'],
        [
            ['push_state(state)', 'スタックに追加', '戦闘開始（フィールドの上に重ねる）'],
            ['pop_state()', '最上位を除去', '戦闘終了（フィールドに戻る）'],
            ['change_state(state)', '最上位を入替', 'タイトル → フィールド'],
        ]
    )

    doc.add_heading('5.2 メインループ', level=2)
    add_code_block(doc,
        'run():\n'
        '  while running:\n'
        '    _handle_events()   # イベント処理 → current_state.handle_event()\n'
        '    _update()          # ステート更新 → current_state.update()\n'
        '    _draw()            # 描画 → current_state.draw()\n'
        '    clock.tick(FPS)    # 60FPS制御'
    )

    doc.add_heading('5.3 ステートインターフェース', level=2)
    doc.add_paragraph('各ステートは以下のメソッドを実装:')
    add_table(doc,
        ['メソッド', '引数', '責務'],
        [
            ['handle_event(event)', 'pygame.event', 'キー入力・イベント処理'],
            ['update()', 'なし', 'フレームごとの状態更新'],
            ['draw(surface)', 'pygame.Surface', '画面描画'],
        ]
    )

    doc.add_page_break()

    # ===== 6. マップシステム =====
    doc.add_heading('6. マップシステム', level=1)

    doc.add_heading('6.1 マップ定義形式', level=2)
    doc.add_paragraph(
        'マップは data.py に文字列リストとして定義。各文字が1タイルに対応。'
    )
    add_code_block(doc,
        'FIELD_MAP = [\n'
        '    "WWWWWWWWWWWWWWWWWWWWWWWWWWWWWWWWWWWWWWWW",  # 40文字 = 40タイル幅\n'
        '    "WGGGGMMMGGGGGGGGGGGGGGGGGGGGGMMMMGGGGGWW",\n'
        '    ...\n'
        ']  # 40行 = 40タイル高さ'
    )

    doc.add_heading('6.2 マップ一覧', level=2)
    add_table(doc,
        ['マップ名', 'サイズ', 'エンカウント率', 'テーブル', '説明'],
        [
            ['field', '40×40', '15% (可変)', '距離依存', 'フィールドマップ'],
            ['castle', '16×16', 'なし', '-', 'ラダトーム城'],
            ['town', '16×16', 'なし', '-', '町'],
            ['dungeon_b1', '16×16', '20%', 'dungeon', 'ダンジョンB1'],
            ['dungeon_b2', '16×16', '25%', 'dungeon', 'ダンジョンB2'],
        ]
    )

    doc.add_heading('6.3 ワープ定義', level=2)
    doc.add_paragraph('マップ間の移動はワープポイントで管理。タイルに到達すると自動的に転送。')
    add_table(doc,
        ['出発マップ', '座標', '到着マップ', '座標'],
        [
            ['field', '(19,14)', 'castle', '(7,14)'],
            ['castle', '(7,14)', 'field', '(19,14)'],
            ['field', '(12,19)', 'town', '(7,15)'],
            ['town', '(7,15)', 'field', '(12,19)'],
            ['field', '(11,6)', 'dungeon_b1', '(2,2)'],
            ['field', '(9,29)', 'dungeon_b1', '(2,2)'],
            ['dungeon_b1', '(2,2)', 'field', '(11,6)'],
            ['dungeon_b1', '(9,13)', 'dungeon_b2', '(2,2)'],
            ['dungeon_b2', '(2,2)', 'dungeon_b1', '(9,13)'],
        ]
    )

    doc.add_heading('6.4 エンカウント判定', level=2)
    doc.add_paragraph(
        'フィールドマップでは城（座標 19,14）からの距離でエンカウントテーブルが変化:'
    )
    add_code_block(doc,
        '距離 < 10  → field_start （低レベル敵）\n'
        '距離 < 20  → field_mid   （中レベル敵）\n'
        '距離 >= 20 → field_far   （高レベル敵）'
    )
    doc.add_paragraph('エンカウント対象タイル: 草、森、砂漠、沼、橋')

    doc.add_heading('6.5 宝箱', level=2)
    add_table(doc,
        ['マップ', '座標', '内容'],
        [
            ['dungeon_b1', '(7, 6)', '120ゴールド'],
            ['dungeon_b2', '(4, 9)', 'てつのおの'],
        ]
    )
    doc.add_paragraph('開封済み宝箱は Player.opened_chests で追跡し、再取得不可。描画時は床タイルに変更。')

    doc.add_heading('6.6 TileMap クラス', level=2)
    add_code_block(doc,
        'class TileMap:\n'
        '    tiles: list[list[int]]   # 2次元タイルID配列\n'
        '\n'
        '    get_tile(x, y) → int           # 座標のタイルID取得（範囲外は-1）\n'
        '    is_passable(x, y) → bool       # 通行可能判定\n'
        '    draw(surface, camera_x, camera_y, opened_chests)  # カメラ基準で描画'
    )
    doc.add_paragraph('カメラシステム: プレイヤーを常に画面中央に配置。表示範囲外のタイルはカリング。')

    doc.add_page_break()

    # ===== 7. プレイヤーシステム =====
    doc.add_heading('7. プレイヤーシステム', level=1)

    doc.add_heading('7.1 Player クラス', level=2)
    add_code_block(doc,
        'class Player:\n'
        '    # 基本情報\n'
        "    name: str = 'ゆうしゃ'\n"
        '    level: int              # 1〜30\n'
        '    exp: int                # 経験値\n'
        '\n'
        '    # ステータス\n'
        '    max_hp, hp: int         # 最大HP / 現在HP\n'
        '    max_mp, mp: int         # 最大MP / 現在MP\n'
        '    strength: int           # ちから\n'
        '    agility: int            # すばやさ\n'
        '    gold: int               # 所持金\n'
        '\n'
        '    # 装備\n'
        '    weapon, armor, shield: str  # 武器・鎧・盾\n'
        '\n'
        '    # インベントリ\n'
        '    items: list[str]        # 所持アイテム（最大8個）\n'
        '\n'
        '    # 位置情報\n'
        '    map_name: str           # 現在のマップ名\n'
        '    tile_x, tile_y: int     # タイル座標\n'
        '    direction: tuple        # 向き (dx, dy)\n'
        '\n'
        '    # フラグ\n'
        '    opened_chests: set      # 開封済み宝箱\n'
        '    flags: set              # ストーリーフラグ'
    )

    doc.add_heading('7.2 戦闘力計算式（DQ1準拠）', level=2)
    add_code_block(doc,
        'こうげき力 = ちから + 武器の攻撃力\n'
        'しゅび力   = すばやさ ÷ 2（切り捨て） + 鎧の守備力 + 盾の守備力'
    )

    doc.add_heading('7.3 レベルアップ', level=2)
    doc.add_paragraph(
        'gain_exp(amount) で経験値を獲得。レベルテーブルの必要経験値に達するとレベルアップ。\n'
        'レベルアップ時: max_hp, max_mp, strength, agility が新レベルの値に更新。'
        '現在のHP/MPは増加分だけ回復。新しい呪文を習得する場合がある。'
    )

    doc.add_heading('7.4 アイテム管理', level=2)
    add_table(doc,
        ['メソッド', '動作'],
        [
            ['add_item(name)', 'アイテム追加（最大8個）'],
            ['remove_item(name)', 'アイテム1個削除'],
            ['use_item(name)', '消耗品使用（やくそう: HP25回復）'],
            ['can_equip(name)', '装備可能判定'],
            ['equip(name)', '装備変更、旧装備を返却'],
        ]
    )

    doc.add_page_break()

    # ===== 8. 戦闘システム =====
    doc.add_heading('8. 戦闘システム', level=1)

    doc.add_heading('8.1 戦闘フェーズ', level=2)
    add_code_block(doc,
        'PHASE_APPEAR → PHASE_COMMAND → PHASE_EXECUTE → PHASE_COMMAND（継続）\n'
        '  (敵出現)      (コマンド選択)    (ターン実行)         ↓（勝敗確定時）\n'
        '                    ↓                              PHASE_RESULT\n'
        '              PHASE_SPELL_SELECT                      ↓\n'
        '              PHASE_ITEM_SELECT                  PHASE_LEVEL_UP（該当時）\n'
        '                                                      ↓\n'
        '                                                  PHASE_END (pop_state)'
    )

    doc.add_heading('8.2 コマンド', level=2)
    add_table(doc,
        ['コマンド', '動作'],
        [
            ['たたかう', '通常攻撃'],
            ['じゅもん', '呪文選択（攻撃/回復/状態異常）'],
            ['どうぐ', 'アイテム使用（消耗品のみ）'],
            ['にげる', '逃走判定'],
        ]
    )

    doc.add_heading('8.3 ダメージ計算式', level=2)
    add_code_block(doc,
        'ダメージ = こうげき力 − しゅび力 ÷ 2 ± ランダム(1〜3)\n'
        '最低ダメージ = 1（0以下の場合）'
    )

    doc.add_heading('8.4 ターン順序', level=2)
    add_code_block(doc,
        '判定値 = すばやさ + ランダム(0〜10)\n'
        '判定値が高い方が先攻'
    )

    doc.add_heading('8.5 逃走判定', level=2)
    add_code_block(doc,
        '成功率 = プレイヤーすばやさ ÷ 敵すばやさ × 0.75\n'
        '失敗時は敵のボーナスターン発生'
    )

    doc.add_heading('8.6 状態異常', level=2)
    add_table(doc,
        ['状態', '効果', '回復条件'],
        [
            ['眠り (sleep)', '行動不能', '毎ターン33%で覚醒'],
            ['沈黙 (silence)', '呪文使用不可', '戦闘中は回復しない'],
        ]
    )

    doc.add_heading('8.7 敵の行動AI', level=2)
    add_code_block(doc,
        '30%の確率で呪文を使用（呪文を持っている場合）\n'
        '  → 攻撃呪文: ダメージを与える\n'
        '  → 状態呪文: プレイヤーに状態異常\n'
        '特殊攻撃を持っている場合:\n'
        '  → fire_breath:        16〜24 ダメージ\n'
        '  → fire_breath_strong: 30〜45 ダメージ\n'
        '上記以外:\n'
        '  → 通常攻撃（ダメージ計算式に従う）'
    )

    doc.add_heading('8.8 勝利・敗北処理', level=2)
    p = doc.add_paragraph()
    run = p.add_run('勝利時:')
    run.bold = True
    doc.add_paragraph('経験値とゴールドを獲得。レベルアップ判定（複数レベル上がる場合あり）。')

    p = doc.add_paragraph()
    run = p.add_run('敗北時:')
    run.bold = True
    doc.add_paragraph(
        'HP → 最大HPの50%で復活。MP → 最大MPの50%で復活。ゴールド → 半減。'
        'ラダトーム城（castle, 7, 14）にテレポート。'
    )

    doc.add_page_break()

    # ===== 9. NPC・ショップ・宿屋システム =====
    doc.add_heading('9. NPC・ショップ・宿屋システム', level=1)

    doc.add_heading('9.1 NPC定義形式', level=2)
    add_code_block(doc,
        '{\n'
        '    "name": "NPC名",\n'
        '    "x": タイルX座標,\n'
        '    "y": タイルY座標,\n'
        '    "color": (R, G, B),        # NPC表示色\n'
        '    "message": ["会話テキスト"],\n'
        '    "shop": [...],             # ショップの場合（省略可）\n'
        '    "inn": True                # 宿屋の場合（省略可）\n'
        '}'
    )

    doc.add_heading('9.2 NPC一覧 — ラダトーム城 (castle)', level=2)
    add_table(doc,
        ['NPC名', '座標', '種別', '内容'],
        [
            ['おうさま', '(7,7)', '会話', '旅立ちの激励'],
            ['へいし', '(3,10)', '会話', '操作説明'],
            ['しょうにん', '(12,10)', 'ショップ', 'たけざお、こんぼう、かわのふく、かわのたて、やくそう、たいまつ'],
        ]
    )

    doc.add_heading('9.3 NPC一覧 — 町 (town)', level=2)
    add_table(doc,
        ['NPC名', '座標', '種別', '内容'],
        [
            ['むらびと', '(3,5)', '会話', '村の情報'],
            ['ろうじん', '(12,5)', '会話', 'ロトの伝説'],
            ['ぶきや', '(3,10)', 'ショップ', 'どうのつるぎ、てつのおの、くさりかたびら、てつのたて'],
            ['やどや', '(12,10)', '宿屋', '6ゴールドで全回復'],
        ]
    )

    doc.add_heading('9.4 ショップシステム', level=2)
    doc.add_paragraph(
        '1. NPCに話しかけるとショップメニュー表示\n'
        '2. アイテム選択 → 所持金チェック → インベントリ空きチェック\n'
        '3. 装備品は自動装備（武器/鎧/盾の判定）\n'
        '4. 購入後メッセージ表示\n'
        '5. 売却機能は未実装'
    )

    doc.add_heading('9.5 宿屋システム', level=2)
    doc.add_paragraph(
        '1. NPCに話しかけると宿泊確認（はい/いいえ）\n'
        '2. 料金: 6ゴールド\n'
        '3. 所持金チェック → HP/MP全回復'
    )

    doc.add_page_break()

    # ===== 10. UIシステム =====
    doc.add_heading('10. UIシステム', level=1)

    doc.add_heading('10.1 フォント管理', level=2)
    doc.add_paragraph(
        '自動検出順序: ipagothic → ipaexgothic → notosanscjkjp → notosansjp → '
        'takao → VL-Gothic → フォールバック（デフォルトフォント）'
    )
    doc.add_paragraph('Font クラス（シングルトン）で同一サイズのフォント再生成を防止。')

    doc.add_heading('10.2 Window クラス（DQ風ウィンドウ）', level=2)
    add_code_block(doc,
        '┌─────────────────┐  ← 白枠線（二重線、3px幅）\n'
        '│                 │\n'
        '│  テキスト領域    │  ← 黒背景、パディング12px\n'
        '│                 │\n'
        '└─────────────────┘'
    )
    add_table(doc,
        ['設定', '値'],
        [
            ['枠線幅', '3px'],
            ['パディング', '12px'],
            ['背景色', '黒 (0,0,0)'],
            ['枠線色', '白 (252,252,252)'],
        ]
    )

    doc.add_heading('10.3 TextRenderer クラス（文字送り）', level=2)
    doc.add_paragraph(
        '1文字ずつ表示（DQスタイル）。表示速度: 2フレーム/文字。'
        '全文表示後にカーソル点滅。Zキーで即時全文表示（スキップ）。'
    )

    doc.add_heading('10.4 Menu クラス（選択メニュー）', level=2)
    doc.add_paragraph(
        '三角カーソル（▶）による選択。上下キーで移動（ラップアラウンド）。'
        '複数列レイアウト対応。Zキーで決定（選択インデックス返却）。'
        'Xキーでキャンセル（-1返却）。'
    )

    doc.add_heading('10.5 StatusWindow / MessageWindow', level=2)
    doc.add_paragraph(
        'StatusWindow: 画面右上にHP/MPを常時表示。\n'
        'MessageWindow: 画面下部にメッセージ表示。キュー管理で複数メッセージを順次表示。'
    )

    doc.add_page_break()

    # ===== 11. 描画システム =====
    doc.add_heading('11. 描画システム', level=1)

    doc.add_heading('11.1 タイル描画', level=2)
    doc.add_paragraph(
        '各タイルは48×48ピクセルのSurfaceとしてプログラム描画。スプライト画像は不使用。'
    )
    add_table(doc,
        ['タイル', '描画内容'],
        [
            ['草', '緑の背景'],
            ['水', '青の背景 + 波線'],
            ['山', '茶の背景 + 三角形 + 雪の頂上'],
            ['森', '濃緑の背景 + 木の円'],
            ['城', '石壁 + 塔 + ドア'],
            ['町', '屋根 + ドア'],
            ['洞窟', '暗い背景 + 楕円'],
            ['階段(下)', '遠近法の下り階段'],
            ['階段(上)', '遠近法の上り階段'],
            ['宝箱', '茶色の箱 + 金のハイライト'],
            ['玉座', '椅子 + 肘掛け'],
            ['橋', '木の板（水上）'],
            ['壁', 'レンガパターン'],
            ['ドア', '木のドア'],
            ['屋根', '三角屋根'],
            ['沼', '暗い円'],
        ]
    )

    doc.add_heading('11.2 キャラクター描画', level=2)
    doc.add_paragraph(
        'プレイヤー: 体（青い矩形）、頭（肌色の円）、髪（円弧）、目（向き応じた位置）、足（茶色の矩形）。\n'
        'NPC: プレイヤーと同構造。体の色はNPCごとに設定。'
    )

    doc.add_heading('11.3 カメラシステム', level=2)
    doc.add_paragraph(
        'プレイヤーを常に画面中央に配置。'
        'カメラ座標 = プレイヤー座標 − 画面中央オフセット。'
        '表示範囲外のタイル・NPCはカリング。'
    )

    doc.add_heading('11.4 移動アニメーション', level=2)
    doc.add_paragraph(
        '1タイル移動 = 8フレーム（48px ÷ 8 = 6px/フレーム）。'
        '移動中はピクセル単位のオフセットで滑らかに描画。'
        'キー押しっぱなしで連続移動。'
    )

    doc.add_page_break()

    # ===== 12. データ定義 =====
    doc.add_heading('12. データ定義', level=1)

    doc.add_heading('12.1 レベルテーブル（抜粋）', level=2)
    add_table(doc,
        ['Lv', '必要経験値', 'HP', 'MP', 'ちから', 'すばやさ'],
        [
            ['1', '0', '15', '0', '4', '4'],
            ['2', '7', '22', '0', '5', '4'],
            ['3', '23', '24', '5', '7', '6'],
            ['5', '110', '35', '12', '12', '8'],
            ['10', '1300', '70', '38', '30', '30'],
            ['15', '4000', '100', '68', '50', '48'],
            ['20', '13000', '140', '100', '72', '72'],
            ['25', '34000', '180', '148', '105', '92'],
            ['30', '65535', '210', '200', '140', '110'],
        ]
    )

    doc.add_heading('12.2 呪文一覧', level=2)
    add_table(doc,
        ['呪文名', '習得Lv', 'MP', '種別', '威力/効果'],
        [
            ['ホイミ', '3', '4', '回復', 'HP24回復'],
            ['ギラ', '4', '2', '攻撃', '10ダメージ'],
            ['ラリホー', '7', '2', '状態', '眠り'],
            ['レミーラ', '9', '3', 'フィールド', '照明'],
            ['マホトーン', '10', '2', '状態', '沈黙'],
            ['リレミト', '13', '6', 'フィールド', '脱出'],
            ['ルーラ', '15', '8', 'フィールド', '帰還'],
            ['トヘロス', '15', '2', 'フィールド', 'エンカウント抑制'],
            ['ベホイミ', '17', '10', '回復', 'HP85回復'],
            ['ベギラマ', '19', '5', '攻撃', '58ダメージ'],
        ]
    )

    doc.add_heading('12.3 武器', level=2)
    add_table(doc,
        ['武器名', '攻撃力', '価格'],
        [
            ['なし', '0', '-'],
            ['たけざお', '2', '10G'],
            ['こんぼう', '4', '60G'],
            ['どうのつるぎ', '10', '180G'],
            ['てつのおの', '15', '560G'],
            ['はがねのつるぎ', '20', '1500G'],
            ['ほのおのつるぎ', '28', '9800G'],
            ['ロトのつるぎ', '40', '- (宝)'],
        ]
    )

    doc.add_heading('12.4 鎧', level=2)
    add_table(doc,
        ['鎧名', '守備力', '価格'],
        [
            ['なし', '0', '-'],
            ['ぬののふく', '2', '20G'],
            ['かわのふく', '4', '70G'],
            ['くさりかたびら', '10', '300G'],
            ['てつのよろい', '16', '800G'],
            ['はがねのよろい', '21', '3000G'],
            ['まほうのよろい', '24', '7700G'],
            ['ロトのよろい', '28', '- (宝)'],
        ]
    )

    doc.add_heading('12.5 盾', level=2)
    add_table(doc,
        ['盾名', '守備力', '価格'],
        [
            ['なし', '0', '-'],
            ['かわのたて', '2', '90G'],
            ['てつのたて', '10', '800G'],
            ['みかがみのたて', '20', '14800G'],
        ]
    )

    doc.add_heading('12.6 敵一覧', level=2)
    add_table(doc,
        ['敵名', 'HP', '攻撃', '守備', '素早さ', '経験値', 'ゴールド', '特殊'],
        [
            ['スライム', '3', '5', '3', '3', '1', '2', '-'],
            ['ドラキー', '6', '9', '6', '6', '2', '3', '-'],
            ['ゴースト', '7', '11', '8', '8', '3', '5', '-'],
            ['まほうつかい', '13', '11', '12', '10', '4', '12', 'ギラ'],
            ['おおさそり', '20', '18', '16', '10', '6', '16', '-'],
            ['リカント', '34', '40', '30', '22', '16', '40', '-'],
            ['ゴーレム', '70', '48', '40', '18', '5', '10', '-'],
            ['しにがみのきし', '50', '52', '35', '36', '28', '70', 'ラリホー'],
            ['ドラゴン', '65', '56', '42', '30', '35', '150', '火の息'],
            ['りゅうおう(人型)', '100', '80', '75', '50', '0', '0', 'ベギラマ/ラリホー'],
            ['りゅうおう(竜)', '130', '120', '100', '60', '0', '0', '強火の息'],
        ]
    )

    doc.add_heading('12.7 エンカウントテーブル', level=2)

    p = doc.add_paragraph()
    run = p.add_run('field_start（城近辺）:')
    run.bold = True
    add_table(doc, ['敵名', '出現率'],
        [['スライム', '60%'], ['ドラキー', '30%'], ['ゴースト', '10%']])

    p = doc.add_paragraph()
    run = p.add_run('field_mid（中距離）:')
    run.bold = True
    add_table(doc, ['敵名', '出現率'],
        [['おおさそり', '40%'], ['まほうつかい', '30%'], ['ゴースト', '20%'], ['ドラキー', '10%']])

    p = doc.add_paragraph()
    run = p.add_run('field_far（遠距離）:')
    run.bold = True
    add_table(doc, ['敵名', '出現率'],
        [['リカント', '30%'], ['しにがみのきし', '30%'], ['おおさそり', '20%'], ['まほうつかい', '20%']])

    p = doc.add_paragraph()
    run = p.add_run('dungeon（ダンジョン共通）:')
    run.bold = True
    add_table(doc, ['敵名', '出現率'],
        [['まほうつかい', '30%'], ['おおさそり', '30%'], ['ゴースト', '20%'], ['しにがみのきし', '20%']])

    doc.add_page_break()

    # ===== 13. 既知の制限・今後の課題 =====
    doc.add_heading('13. 既知の制限・今後の課題', level=1)
    add_table(doc,
        ['項目', '状態', '説明'],
        [
            ['セーブ/ロード', '未実装', 'ゲーム進行の保存・読み込み'],
            ['音楽・効果音', '未実装', 'BGM、SE の再生'],
            ['フィールド呪文', '簡略化', '最初の回復呪文のみフィールドで使用可能'],
            ['ショップ売却', '未実装', 'アイテムの売却機能'],
            ['ダンジョン照明', '未実装', '暗闘・たいまつ/レミーラによる照明'],
            ['竜王の城', '未実装', '最終ダンジョンの実装'],
            ['NPC アニメーション', '未実装', 'NPCの移動・方向変化'],
            ['戦闘アニメーション', '未実装', '攻撃・呪文のエフェクト'],
            ['ストーリー進行', '未実装', 'クエストフラグによるストーリー管理'],
            ['スプライト対応', '未実装', '画像アセットによるキャラクター・タイル描画'],
        ]
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('本設計書は DQ1-Sample リポジトリのソースコードに基づいて作成されました。')
    run.italic = True

    # 保存
    output_path = '/home/user/DQ1-Sample/DQ1-Sample_設計書.docx'
    doc.save(output_path)
    print(f'設計書を出力しました: {output_path}')


if __name__ == '__main__':
    main()
